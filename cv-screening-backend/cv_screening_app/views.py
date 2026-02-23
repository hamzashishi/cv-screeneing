from rest_framework import viewsets, status, permissions
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework.parsers import MultiPartParser, FormParser, JSONParser
from django.conf import settings
from django.shortcuts import get_object_or_404
from django.db import IntegrityError, transaction
from django.utils import timezone
from django.core.mail import send_mail
from django.core.files.base import ContentFile
from django.http import HttpResponse
from html import escape
from docx import Document
import io
import json
import re
from datetime import timedelta

from cv_screening_app.models import (
    CustomUser, HRPersonnel, Applicant, JobPosting, CVUpload,
    JobApplication, ParsedCVData, Notification, ScreeningCriteria, EmailVerificationOTP, PasswordResetOTP
)
from cv_screening_app.serializers import (
    CustomUserSerializer, CustomUserCreateSerializer, HRPersonnelSerializer,
    ApplicantSerializer, JobPostingSerializer, CVUploadSerializer,
    JobApplicationSerializer, NotificationSerializer, ScreeningCriteriaSerializer
)
from cv_screening_app.cv_parser import CVParser
from cv_screening_app.ranking_engine import RankingEngine
from cv_screening_app.authentication import generate_jwt_token


def ensure_parsed_data_record(cv_upload):
    """Ensure a CVUpload has ParsedCVData; parse file on demand if missing."""
    if not cv_upload:
        return None
    try:
        existing_record = cv_upload.parsed_data_record
        has_summary = bool(existing_record.summary and str(existing_record.summary).strip())
        has_identity = bool(
            (existing_record.full_name and str(existing_record.full_name).strip()) or
            (existing_record.email and str(existing_record.email).strip()) or
            (existing_record.phone_number and str(existing_record.phone_number).strip()) or
            (existing_record.location and str(existing_record.location).strip())
        )
        has_core_nlp = bool(
            (existing_record.skills and len(existing_record.skills) > 0) or
            (existing_record.education and len(existing_record.education) > 0) or
            (existing_record.work_experience and len(existing_record.work_experience) > 0) or
            (existing_record.total_years_experience is not None)
        )
        if has_summary and (has_identity or has_core_nlp):
            return existing_record
        # Re-parse if summary exists but critical extracted fields are still missing.
    except ParsedCVData.DoesNotExist:
        existing_record = None

    parser = CVParser()
    parsed_data = None
    try:
        parsed_data = parser.parse_cv(cv_upload.cv_file.path)
    except Exception:
        parsed_data = None

    if parsed_data:
        cv_upload.extracted_text = parsed_data.get('extracted_text', '') or ''
        cv_upload.save(update_fields=['extracted_text'])
        record, _ = ParsedCVData.objects.update_or_create(
            cv_upload=cv_upload,
            defaults={
                'full_name': parsed_data.get('full_name'),
                'email': parsed_data.get('email'),
                'phone_number': parsed_data.get('phone_number'),
                'location': parsed_data.get('location'),
                'summary': parsed_data.get('summary'),
                'skills': parsed_data.get('skills', []),
                'education': parsed_data.get('education', []),
                'highest_education': parsed_data.get('highest_education') or 'high_school',
                'work_experience': parsed_data.get('work_experience', []),
                'total_years_experience': parsed_data.get('total_years_experience') or 0,
                'certifications': parsed_data.get('certifications', []),
                'languages': parsed_data.get('languages', []),
                'extraction_confidence': 100,
            }
        )
        return record

    record, _ = ParsedCVData.objects.update_or_create(
        cv_upload=cv_upload,
        defaults={
            'full_name': None,
            'email': None,
            'phone_number': None,
            'location': None,
            'summary': '',
            'skills': [],
            'education': [],
            'highest_education': 'high_school',
            'work_experience': [],
            'total_years_experience': 0,
            'certifications': [],
            'languages': [],
            'extraction_confidence': 0,
        }
    )
    return record


class IsHRUser(permissions.BasePermission):
    def has_permission(self, request, view):
        return request.user and request.user.role == 'hr'


class IsApplicantUser(permissions.BasePermission):
    def has_permission(self, request, view):
        return request.user and request.user.role == 'applicant'


class IsAdminUser(permissions.BasePermission):
    def has_permission(self, request, view):
        return bool(request.user and request.user.is_authenticated and (request.user.is_staff or request.user.is_superuser))


class UserAuthViewSet(viewsets.ViewSet):
    """User registration and authentication endpoints"""
    permission_classes = [permissions.AllowAny]
    parser_classes = (MultiPartParser, FormParser, JSONParser)

    def _send_verification_otp(self, user):
        EmailVerificationOTP.objects.filter(user=user, is_used=False).update(is_used=True)
        otp_code = EmailVerificationOTP.generate_code()
        expires_at = timezone.now() + timedelta(minutes=10)
        otp = EmailVerificationOTP.objects.create(
            user=user,
            otp_code=otp_code,
            expires_at=expires_at,
        )

        subject = 'Your CV Screening verification code'
        message = (
            f"Hello {user.first_name or user.username},\n\n"
            f"Your OTP verification code is: {otp_code}\n"
            f"This code expires in 10 minutes.\n\n"
            f"If you did not request this, please ignore this email."
        )
        from_email = getattr(settings, 'DEFAULT_FROM_EMAIL', 'no-reply@cvscreening.local')
        send_mail(subject, message, from_email, [user.email], fail_silently=False)
        return otp

    def _send_password_reset_otp(self, user):
        PasswordResetOTP.objects.filter(user=user, is_used=False).update(is_used=True)
        otp_code = PasswordResetOTP.generate_code()
        expires_at = timezone.now() + timedelta(minutes=10)
        otp = PasswordResetOTP.objects.create(
            user=user,
            otp_code=otp_code,
            expires_at=expires_at,
        )

        subject = 'Your CV Screening password reset code'
        message = (
            f"Hello {user.first_name or user.username},\n\n"
            f"Your password reset OTP code is: {otp_code}\n"
            f"This code expires in 10 minutes.\n\n"
            f"If you did not request this, please ignore this email."
        )
        from_email = getattr(settings, 'DEFAULT_FROM_EMAIL', 'no-reply@cvscreening.local')
        sent = True
        try:
            send_mail(subject, message, from_email, [user.email], fail_silently=False)
        except Exception:
            sent = False
        return otp, sent
    
    def register(self, request):
        """Register new user"""
    @action(detail=False, methods=['get', 'post'])
    def register(self, request):
        """Register new user or return usage info for GET"""
        if request.method == 'GET':
            return Response({
                'info': 'POST JSON to register. Required: email, username, password, password_confirm, role'
            }, status=status.HTTP_200_OK)
        import logging
        logger = logging.getLogger(__name__)

        serializer = CustomUserCreateSerializer(data=request.data)
        if serializer.is_valid():
            try:
                user = serializer.save()
            except Exception as exc:
                logger.exception('Error creating user: %s', exc)
                return Response({'error': 'Failed to create user'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            # Create role-specific profile
            if user.role == 'hr':
                HRPersonnel.objects.create(
                    user=user,
                    company_name=request.data.get('company_name', ''),
                    company_description=request.data.get('company_description', ''),
                    company_location=request.data.get('company_location', '')
                )
            elif user.role == 'applicant':
                Applicant.objects.create(
                    user=user,
                    location=request.data.get('location', '')
                )

            user.is_active = False
            user.is_active_user = False
            user.save(update_fields=['is_active', 'is_active_user'])

            try:
                self._send_verification_otp(user)
            except Exception:
                return Response(
                    {'error': 'Registration created but failed to send OTP. Please try resend OTP.'},
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR
                )

            response_data = {
                'message': 'Registration successful. Please verify OTP sent to your email.',
                'requires_email_verification': True,
                'pending_admin_approval': user.role == 'hr',
                'email': user.email,
                'user': CustomUserSerializer(user).data,
            }
            return Response(response_data, status=status.HTTP_201_CREATED)
        
        # Log serializer errors for easier debugging
        import logging
        logging.getLogger(__name__).warning('Registration failed: %s', serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
    
    def login(self, request):
        """User login endpoint"""
    @action(detail=False, methods=['get', 'post'])
    def login(self, request):
        """User login or return usage info for GET"""
        if request.method == 'GET':
            return Response({
                'info': 'POST JSON to login. Required: email and password'
            }, status=status.HTTP_200_OK)
        email = request.data.get('email')
        password = request.data.get('password')
        
        if not email or not password:
            return Response(
                {'error': 'Email and password are required'},
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            # Login by email only (case-insensitive)
            user = CustomUser.objects.filter(email__iexact=email).first()

            if user and user.check_password(password):
                if not user.is_active or not user.is_active_user:
                    return Response({'error': 'Account is not active. Please verify OTP first.'}, status=status.HTTP_403_FORBIDDEN)

                if user.role == 'hr' and not (user.is_staff or user.is_superuser):
                    hr_profile = HRPersonnel.objects.filter(user=user).first()
                    if not hr_profile:
                        return Response({'error': 'HR profile not found'}, status=status.HTTP_403_FORBIDDEN)
                    if not hr_profile.is_verified:
                        return Response(
                            {'error': 'Your HR account is pending admin approval. Please wait for administrator verification.'},
                            status=status.HTTP_403_FORBIDDEN
                        )

                token = generate_jwt_token(user)
                return Response({
                    'message': 'Login successful',
                    'token': token,
                    'user': CustomUserSerializer(user).data
                }, status=status.HTTP_200_OK)

            if user:
                return Response({'error': 'Invalid credentials'}, status=status.HTTP_401_UNAUTHORIZED)

            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as exc:
            return Response({'error': str(exc)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    @action(detail=False, methods=['post'])
    def verify_otp(self, request):
        """Verify email OTP and activate account."""
        email = (request.data.get('email') or '').strip().lower()
        otp_code = (request.data.get('otp') or '').strip()

        if not email or not otp_code:
            return Response({'error': 'Email and OTP are required'}, status=status.HTTP_400_BAD_REQUEST)

        user = CustomUser.objects.filter(email__iexact=email).first()
        if not user:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

        otp = EmailVerificationOTP.objects.filter(
            user=user,
            otp_code=otp_code,
            is_used=False,
            expires_at__gt=timezone.now(),
        ).order_by('-created_at').first()

        if not otp:
            return Response({'error': 'Invalid or expired OTP'}, status=status.HTTP_400_BAD_REQUEST)

        otp.is_used = True
        otp.save(update_fields=['is_used'])
        user.is_active = True
        user.is_active_user = True
        user.save(update_fields=['is_active', 'is_active_user'])

        token = generate_jwt_token(user)
        return Response({
            'message': 'Email verified successfully',
            'token': token,
            'user': CustomUserSerializer(user).data
        }, status=status.HTTP_200_OK)

    @action(detail=False, methods=['post'])
    def resend_otp(self, request):
        """Resend OTP for unverified account."""
        email = (request.data.get('email') or '').strip().lower()
        if not email:
            return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

        user = CustomUser.objects.filter(email__iexact=email).first()
        if not user:
            return Response({'error': 'User not found'}, status=status.HTTP_404_NOT_FOUND)

        if user.is_active:
            return Response({'error': 'Email already verified'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            self._send_verification_otp(user)
        except Exception:
            return Response({'error': 'Failed to resend OTP'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        response_data = {'message': 'OTP resent successfully'}
        return Response(response_data, status=status.HTTP_200_OK)

    @action(detail=False, methods=['post'])
    def forgot_password(self, request):
        """Send password reset OTP to email."""
        email = (request.data.get('email') or '').strip().lower()
        if not email:
            return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

        user = CustomUser.objects.filter(email__iexact=email).first()
        if not user:
            return Response(
                {'message': 'If this email exists, a password reset code has been sent.'},
                status=status.HTTP_200_OK
            )

        otp, sent = self._send_password_reset_otp(user)

        response = {
            'message': 'If this email exists, a password reset code has been sent.'
            if sent else
            'Email delivery failed. Please try again later or contact support.'
        }
        return Response(response, status=status.HTTP_200_OK if sent else status.HTTP_503_SERVICE_UNAVAILABLE)

    @action(detail=False, methods=['post'])
    def reset_password(self, request):
        """Reset password using email + OTP."""
        email = (request.data.get('email') or '').strip().lower()
        otp_code = (request.data.get('otp') or '').strip()
        new_password = request.data.get('new_password') or ''
        confirm_password = request.data.get('confirm_password') or ''

        if not email or not otp_code or not new_password or not confirm_password:
            return Response(
                {'error': 'Email, OTP, new password, and confirm password are required'},
                status=status.HTTP_400_BAD_REQUEST
            )

        if new_password != confirm_password:
            return Response({'error': 'Passwords do not match'}, status=status.HTTP_400_BAD_REQUEST)

        if len(new_password) < 8:
            return Response({'error': 'Password must be at least 8 characters'}, status=status.HTTP_400_BAD_REQUEST)

        user = CustomUser.objects.filter(email__iexact=email).first()
        if not user:
            return Response({'error': 'Invalid reset request'}, status=status.HTTP_400_BAD_REQUEST)

        otp = PasswordResetOTP.objects.filter(
            user=user,
            otp_code=otp_code,
            is_used=False,
            expires_at__gt=timezone.now(),
        ).order_by('-created_at').first()

        if not otp:
            return Response({'error': 'Invalid or expired OTP'}, status=status.HTTP_400_BAD_REQUEST)

        user.set_password(new_password)
        user.save(update_fields=['password', 'updated_at'])
        otp.is_used = True
        otp.save(update_fields=['is_used'])

        return Response({'message': 'Password reset successful. Please login.'}, status=status.HTTP_200_OK)

    @action(detail=False, methods=['get', 'patch'], permission_classes=[permissions.IsAuthenticated])
    def profile(self, request):
        """Get or partially update current authenticated user profile"""
        if request.method == 'GET':
            return Response(CustomUserSerializer(request.user).data, status=status.HTTP_200_OK)

        allowed_fields = {'first_name', 'last_name', 'phone_number', 'profile_picture'}
        update_data = {k: v for k, v in request.data.items() if k in allowed_fields}
        if not update_data:
            return Response({'error': 'No updatable fields provided'}, status=status.HTTP_400_BAD_REQUEST)

        serializer = CustomUserSerializer(request.user, data=update_data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_200_OK)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class HRPersonnelViewSet(viewsets.ModelViewSet):
    """HR Personnel profile management"""
    queryset = HRPersonnel.objects.all()
    serializer_class = HRPersonnelSerializer
    permission_classes = [permissions.IsAuthenticated, IsHRUser]
    
    @action(detail=False, methods=['get'])
    def my_profile(self, request):
        """Get current HR user's profile"""
        hr_profile = get_object_or_404(HRPersonnel, user=request.user)
        serializer = self.get_serializer(hr_profile)
        return Response(serializer.data)


class AdminApprovalViewSet(viewsets.ReadOnlyModelViewSet):
    """Admin-only endpoints to review and approve HR company registrations."""
    queryset = HRPersonnel.objects.select_related('user').all()
    serializer_class = HRPersonnelSerializer
    permission_classes = [permissions.IsAuthenticated, IsAdminUser]

    def get_queryset(self):
        queryset = super().get_queryset()
        status_filter = (self.request.query_params.get('status') or 'pending').strip().lower()
        if status_filter == 'approved':
            return queryset.filter(is_verified=True)
        if status_filter == 'all':
            return queryset
        return queryset.filter(is_verified=False)

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        profile = self.get_object()
        profile.is_verified = True
        profile.save(update_fields=['is_verified', 'updated_at'])
        return Response({'message': 'HR company approved successfully'}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        profile = self.get_object()
        profile.is_verified = False
        profile.user.is_active_user = False
        profile.user.is_active = False
        profile.user.save(update_fields=['is_active_user', 'is_active', 'updated_at'])
        return Response({'message': 'HR account rejected and disabled'}, status=status.HTTP_200_OK)


class ApplicantViewSet(viewsets.ModelViewSet):
    """Applicant profile management"""
    queryset = Applicant.objects.all()
    serializer_class = ApplicantSerializer
    permission_classes = [permissions.IsAuthenticated, IsApplicantUser]
    
    @action(detail=False, methods=['get'])
    def my_profile(self, request):
        """Get current applicant's profile"""
        applicant_profile = get_object_or_404(Applicant, user=request.user)
        serializer = self.get_serializer(applicant_profile)
        return Response(serializer.data)


class JobPostingViewSet(viewsets.ModelViewSet):
    """Job posting management"""
    queryset = JobPosting.objects.filter(is_active=True)
    serializer_class = JobPostingSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        if self.request.user.role == 'hr':
            return JobPosting.objects.filter(hr__user=self.request.user)
        return JobPosting.objects.filter(is_active=True)
    
    def perform_create(self, serializer):
        hr_profile = get_object_or_404(HRPersonnel, user=self.request.user)
        serializer.save(hr=hr_profile)
        
        # Create default screening criteria
        job = serializer.instance
        ScreeningCriteria.objects.create(job=job)

    def create(self, request, *args, **kwargs):
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can create jobs'}, status=status.HTTP_403_FORBIDDEN)
        return super().create(request, *args, **kwargs)

    def _get_hr_job_or_forbidden(self, pk=None):
        if self.request.user.role != 'hr':
            return None, Response({'error': 'Only HR users can modify jobs'}, status=status.HTTP_403_FORBIDDEN)
        job = get_object_or_404(JobPosting, pk=pk)
        if job.hr.user_id != self.request.user.id:
            return None, Response({'error': 'Not allowed to modify this job'}, status=status.HTTP_403_FORBIDDEN)
        return job, None

    def update(self, request, *args, **kwargs):
        _, forbidden = self._get_hr_job_or_forbidden(kwargs.get('pk'))
        if forbidden:
            return forbidden
        return super().update(request, *args, **kwargs)

    def partial_update(self, request, *args, **kwargs):
        _, forbidden = self._get_hr_job_or_forbidden(kwargs.get('pk'))
        if forbidden:
            return forbidden
        return super().partial_update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        job, forbidden = self._get_hr_job_or_forbidden(kwargs.get('pk'))
        if forbidden:
            return forbidden
        # Soft-delete to preserve application history.
        job.is_active = False
        job.save(update_fields=['is_active'])
        return Response({'message': 'Job deleted successfully'}, status=status.HTTP_200_OK)
    
    @action(detail=True, methods=['get'])
    def applicants(self, request, pk=None):
        """Get applicants for a job"""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can view applicants'}, status=status.HTTP_403_FORBIDDEN)
        job = self.get_object()
        if job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to view applicants for this job'}, status=status.HTTP_403_FORBIDDEN)
        applications = JobApplication.objects.filter(job=job).select_related(
            'applicant__user', 'cv', 'cv__parsed_data_record'
        ).order_by('-application_score')
        for application in applications:
            cv = getattr(application, 'cv', None)
            if cv:
                ensure_parsed_data_record(cv)
        serializer = JobApplicationSerializer(applications, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get'])
    def export_applicants_excel(self, request, pk=None):
        """Export applicants with details as Excel-compatible XML (.xls)."""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can export applicants'}, status=status.HTTP_403_FORBIDDEN)
        job = self.get_object()
        if job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to export applicants for this job'}, status=status.HTTP_403_FORBIDDEN)
        applications = JobApplication.objects.filter(job=job).select_related(
            'applicant__user', 'cv', 'cv__parsed_data_record'
        ).order_by('-application_score', '-applied_date')

        headers = [
            'Applicant Name',
            'Email',
            'Phone',
            'Location',
            'Status',
            'Rank',
            'Overall Score',
            'Skills Score',
            'Experience Score',
            'Education Score',
            'Total Years Experience',
            'Highest Education',
            'Skills',
            'Work Experience',
            'Training/Workshop',
            'Languages',
            'Certifications',
            'Applied Date',
            'CV Type',
            'CV URL',
            'HR Notes',
        ]

        def format_experience(items):
            lines = []
            for item in items or []:
                if isinstance(item, dict):
                    org = str(item.get('organization') or item.get('institution') or '').strip()
                    role = str(item.get('position') or '').strip()
                    date_range = str(item.get('date_range') or '').strip()
                    duties = str(item.get('responsibilities') or item.get('description') or item.get('duties') or '').strip()
                    head = ' '.join(part for part in [org, role] if part)
                    if date_range:
                        head = f"{head} ({date_range})" if head else f"({date_range})"
                    if duties:
                        line = f"{head} - {duties}" if head else duties
                    else:
                        line = head
                    if line:
                        lines.append(line)
                else:
                    line = str(item).strip()
                    if line:
                        lines.append(line)
            return ' | '.join(lines)

        rows = []
        for app in applications:
            cv_url = ''
            cv_type = ''
            parsed = None
            if app.cv and getattr(app.cv, 'cv_file', None):
                try:
                    cv_url = request.build_absolute_uri(app.cv.cv_file.url)
                except Exception:
                    cv_url = ''
            if app.cv:
                parsed = ensure_parsed_data_record(app.cv)
                parsed_meta = getattr(app.cv, 'parsed_data', None) or {}
                cv_type = 'Template' if parsed_meta.get('template_source') == 'builder' else 'Uploaded'

            skills = (parsed.skills or []) if parsed else []
            education = (parsed.education or []) if parsed else []
            work_experience = (parsed.work_experience or []) if parsed else []
            certifications = (parsed.certifications or []) if parsed else []
            languages = (parsed.languages or []) if parsed else []
            total_years = parsed.total_years_experience if parsed else ''
            highest_education = parsed.highest_education if parsed else ''
            training_workshop = ''
            if app.cv:
                parsed_meta = getattr(app.cv, 'parsed_data', None) or {}
                training_workshop = format_experience(parsed_meta.get('training_workshop') or [])

            applicant_name = app.applicant.user.get_full_name() or ''
            if parsed and parsed.full_name:
                applicant_name = parsed.full_name

            rows.append([
                applicant_name,
                (parsed.email if parsed else None) or app.applicant.user.email or '',
                (parsed.phone_number if parsed else None) or app.applicant.user.phone_number or '',
                (parsed.location if parsed else None) or app.applicant.location or '',
                app.status or '',
                str(app.application_rank or ''),
                str(round(float(app.application_score), 2)) if app.application_score is not None else '',
                str(round(float(app.skills_score), 2)) if app.skills_score is not None else '',
                str(round(float(app.experience_score), 2)) if app.experience_score is not None else '',
                str(round(float(app.education_score), 2)) if app.education_score is not None else '',
                str(total_years) if total_years != '' else '',
                str(highest_education or ''),
                ', '.join(skills[:20]),
                format_experience(work_experience),
                training_workshop,
                ', '.join(languages[:10]),
                ', '.join(certifications[:10]),
                app.applied_date.strftime('%Y-%m-%d %H:%M') if app.applied_date else '',
                cv_type,
                cv_url,
                app.hr_notes or '',
            ])

        def make_cell(value: str) -> str:
            return f'<Cell><Data ss:Type="String">{escape(str(value or ""))}</Data></Cell>'

        xml_rows = []
        xml_rows.append('<Row>' + ''.join(make_cell(h) for h in headers) + '</Row>')
        for row in rows:
            xml_rows.append('<Row>' + ''.join(make_cell(col) for col in row) + '</Row>')

        worksheet_name = f"{job.job_title[:24]} Applicants" if job.job_title else "Applicants"
        xml_content = f'''<?xml version="1.0"?>
<Workbook xmlns="urn:schemas-microsoft-com:office:spreadsheet"
 xmlns:o="urn:schemas-microsoft-com:office:office"
 xmlns:x="urn:schemas-microsoft-com:office:excel"
 xmlns:ss="urn:schemas-microsoft-com:office:spreadsheet"
 xmlns:html="http://www.w3.org/TR/REC-html40">
 <Worksheet ss:Name="{escape(worksheet_name)}">
  <Table>
   {''.join(xml_rows)}
  </Table>
 </Worksheet>
</Workbook>'''

        safe_job_title = ''.join(ch if ch.isalnum() or ch in ('_', '-') else '_' for ch in (job.job_title or 'job').replace(' ', '_'))
        stamp = timezone.now().strftime('%Y%m%d_%H%M')
        filename = f'{safe_job_title}_applicants_{stamp}.xls'
        response = HttpResponse(xml_content, content_type='application/vnd.ms-excel')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    @action(detail=True, methods=['post'])
    def rank_candidates(self, request, pk=None):
        """Trigger candidate ranking for a job"""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can rank candidates'}, status=status.HTTP_403_FORBIDDEN)
        job = self.get_object()
        if job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to rank candidates for this job'}, status=status.HTTP_403_FORBIDDEN)
        applications = JobApplication.objects.filter(job=job)
        
        try:
            ranking_engine = RankingEngine({
                'skills_weight': job.screening_criteria.skills_weight,
                'experience_weight': job.screening_criteria.experience_weight,
                'education_weight': job.screening_criteria.education_weight,
            })
            
            candidates_data = []
            app_map = {}
            
            for app in applications:
                if not app.cv:
                    continue
                parsed = ensure_parsed_data_record(app.cv)
                if not parsed:
                    continue
                candidates_data.append({
                    'id': str(app.id),
                    'skills': parsed.skills,
                    'highest_education': parsed.highest_education,
                    'total_years_experience': parsed.total_years_experience,
                    'extracted_text': app.cv.extracted_text or '',
                })
                app_map[str(app.id)] = app
            
            if not candidates_data:
                return Response({'error': 'No candidates with parsed data'}, status=status.HTTP_400_BAD_REQUEST)

            # Reset previous ranking output for this job before writing fresh ranks.
            applications.update(
                application_score=None,
                application_rank=None,
                skills_score=None,
                experience_score=None,
                education_score=None,
            )
            
            # Get job requirements
            job_data = {
                'required_skills': job.required_skills,
                'required_education': job.required_education,
                'years_of_experience_min': job.years_of_experience_min,
                'years_of_experience_max': job.years_of_experience_max,
            }
            
            ranked = ranking_engine.rank_candidates(job_data, candidates_data)
            
            # Update scores in database
            for rank, (candidate, score) in enumerate(ranked, 1):
                app = app_map[candidate['id']]
                app.application_score = score
                app.application_rank = rank
                app.skills_score = candidate['scores']['skills_score']
                app.experience_score = candidate['scores']['experience_score']
                app.education_score = candidate['scores']['education_score']
                app.save()
            
            return Response({'message': 'Candidates ranked successfully'}, status=status.HTTP_200_OK)
        
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['get', 'patch', 'post'])
    def screening_criteria(self, request, pk=None):
        """Get or update screening criteria for a job"""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can manage screening criteria'}, status=status.HTTP_403_FORBIDDEN)
        job = self.get_object()
        if job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to manage this job criteria'}, status=status.HTTP_403_FORBIDDEN)
        try:
            criteria = job.screening_criteria
        except ScreeningCriteria.DoesNotExist:
            criteria = ScreeningCriteria.objects.create(job=job)

        if request.method == 'GET':
            serializer = ScreeningCriteriaSerializer(criteria)
            return Response(serializer.data)

        # Update
        serializer = ScreeningCriteriaSerializer(criteria, data=request.data, partial=True)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class CVUploadViewSet(viewsets.ModelViewSet):
    """CV upload and parsing"""
    queryset = CVUpload.objects.all()
    serializer_class = CVUploadSerializer
    permission_classes = [permissions.IsAuthenticated, IsApplicantUser]
    parser_classes = (MultiPartParser, FormParser, JSONParser)
    
    def get_queryset(self):
        applicant = get_object_or_404(Applicant, user=self.request.user)
        return CVUpload.objects.filter(applicant=applicant)
    
    def perform_create(self, serializer):
        applicant = get_object_or_404(Applicant, user=self.request.user)
        cv_upload = serializer.save(applicant=applicant)
        
        # Parse CV asynchronously or synchronously
        try:
            parser = CVParser()
            parsed_data = parser.parse_cv(cv_upload.cv_file.path)
            
            if parsed_data:
                cv_upload.extracted_text = parsed_data.get('extracted_text', '')
                cv_upload.save()
                
                # Save parsed data
                ParsedCVData.objects.update_or_create(
                    cv_upload=cv_upload,
                    defaults={
                        'full_name': parsed_data.get('full_name'),
                        'email': parsed_data.get('email'),
                        'phone_number': parsed_data.get('phone_number'),
                        'location': parsed_data.get('location'),
                        'summary': parsed_data.get('summary'),
                        'skills': parsed_data.get('skills', []),
                        'education': parsed_data.get('education', []),
                        'highest_education': parsed_data.get('highest_education'),
                        'work_experience': parsed_data.get('work_experience', []),
                        'total_years_experience': parsed_data.get('total_years_experience'),
                        'certifications': parsed_data.get('certifications', []),
                        'languages': parsed_data.get('languages', []),
                    }
                )
            else:
                # Ensure there is always a parsed_data_record for uploaded PDF/DOCX.
                ParsedCVData.objects.update_or_create(
                    cv_upload=cv_upload,
                    defaults={
                        'full_name': None,
                        'email': None,
                        'phone_number': None,
                        'location': None,
                        'summary': '',
                        'skills': [],
                        'education': [],
                        'highest_education': 'high_school',
                        'work_experience': [],
                        'total_years_experience': 0,
                        'certifications': [],
                        'languages': [],
                        'extraction_confidence': 0,
                    }
                )
        except Exception as e:
            print(f"Error parsing CV: {str(e)}")
            ParsedCVData.objects.update_or_create(
                cv_upload=cv_upload,
                defaults={
                    'full_name': None,
                    'email': None,
                    'phone_number': None,
                    'location': None,
                    'summary': '',
                    'skills': [],
                    'education': [],
                    'highest_education': 'high_school',
                    'work_experience': [],
                    'total_years_experience': 0,
                    'certifications': [],
                    'languages': [],
                    'extraction_confidence': 0,
                }
            )
    
    @action(detail=True, methods=['post'])
    def set_primary(self, request, pk=None):
        """Set CV as primary for screening"""
        cv = self.get_object()
        CVUpload.objects.filter(applicant=cv.applicant).update(is_primary=False)
        cv.is_primary = True
        cv.save()
        return Response({'message': 'CV set as primary'}, status=status.HTTP_200_OK)

    @action(detail=False, methods=['post'])
    def create_from_template(self, request):
        """Create a CV from structured template data and store parsed fields for ranking."""
        applicant = get_object_or_404(Applicant, user=request.user)

        def parse_list_field(value):
            if value is None:
                return []
            if isinstance(value, list):
                return value
            if isinstance(value, str):
                raw = value.strip()
                if not raw:
                    return []
                try:
                    parsed = json.loads(raw)
                    if isinstance(parsed, list):
                        return parsed
                except Exception:
                    pass
                return [v.strip() for v in raw.split(',') if v.strip()]
            return []

        def parse_date_range_line(value):
            if not value:
                return None
            raw = str(value).strip()
            if not raw:
                return None
            clean = re.sub(r'^[\(\[]\s*|\s*[\)\]]$', '', raw).strip()

            month_map = {
                'jan': 1, 'january': 1,
                'feb': 2, 'february': 2,
                'mar': 3, 'march': 3,
                'apr': 4, 'april': 4,
                'may': 5,
                'jun': 6, 'june': 6,
                'jul': 7, 'july': 7,
                'aug': 8, 'august': 8,
                'sep': 9, 'sept': 9, 'september': 9,
                'oct': 10, 'october': 10,
                'nov': 11, 'november': 11,
                'dec': 12, 'december': 12,
            }
            month_regex = r'(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t)?(?:ember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)'
            dash = r'[-\u2013\u2014]'

            pattern_full = rf'(?P<sm>{month_regex})\s+(?P<sy>\d{{4}})\s*{dash}\s*(?P<em>{month_regex})\s+(?P<ey>\d{{4}}|present|current)'
            match = re.search(pattern_full, clean, re.IGNORECASE)
            if match:
                sm = month_map[match.group('sm').lower()]
                sy = int(match.group('sy'))
                em = month_map[match.group('em').lower()]
                ey_raw = match.group('ey').lower()
                ey = timezone.now().year if ey_raw in ('present', 'current') else int(ey_raw)
                return (sy, sm, ey, em)

            pattern_same_year = rf'(?P<sm>{month_regex})\s*{dash}\s*(?P<em>{month_regex})\s+(?P<ey>\d{{4}})'
            match = re.search(pattern_same_year, clean, re.IGNORECASE)
            if match:
                sm = month_map[match.group('sm').lower()]
                em = month_map[match.group('em').lower()]
                ey = int(match.group('ey'))
                return (ey, sm, ey, em)

            pattern_years = rf'(?P<sy>\d{{4}})\s*{dash}\s*(?P<ey>\d{{4}}|present|current)'
            match = re.search(pattern_years, clean, re.IGNORECASE)
            if match:
                sy = int(match.group('sy'))
                ey_raw = match.group('ey').lower()
                ey = timezone.now().year if ey_raw in ('present', 'current') else int(ey_raw)
                return (sy, 1, ey, 12)

            return None

        def format_work_experience_item(item):
            if isinstance(item, dict):
                org = (item.get('organization') or item.get('company') or item.get('institution') or '').strip()
                date_range = (item.get('date_range') or '').strip()
                responsibilities = (item.get('responsibilities') or item.get('description') or item.get('duties') or '').strip()
                parts = []
                if org:
                    parts.append(org)
                if date_range:
                    parts.append(f"({date_range})")
                line = ' '.join(parts).strip()
                if responsibilities:
                    line = f"{line} - {responsibilities}" if line else responsibilities
                return line or str(item)
            return str(item)

        surname = request.data.get('surname') or ''
        given_names = request.data.get('given_names') or ''
        full_name = request.data.get('full_name') or f"{surname} {given_names}".strip() or request.user.get_full_name()
        sex = request.data.get('sex') or ''
        marital_status = request.data.get('marital_status') or ''
        date_of_birth = request.data.get('date_of_birth') or ''
        nationality = request.data.get('nationality') or ''
        email = request.data.get('email') or request.user.email
        phone_number = request.data.get('phone_number') or request.user.phone_number
        alt_phone = request.data.get('alt_phone') or ''
        location = request.data.get('location', '')
        summary = request.data.get('summary', '')
        highest_education = request.data.get('highest_education', 'bachelor')
        years_experience_raw = request.data.get('total_years_experience', 0)

        try:
            total_years_experience = float(years_experience_raw or 0)
        except (TypeError, ValueError):
            total_years_experience = 0.0

        skills = parse_list_field(request.data.get('skills'))
        education = parse_list_field(request.data.get('education'))
        work_experience = parse_list_field(request.data.get('work_experience'))
        training_workshop = parse_list_field(request.data.get('training_workshop'))
        certifications = parse_list_field(request.data.get('certifications'))
        languages = parse_list_field(request.data.get('languages'))
        hobbies = request.data.get('hobbies', '')
        referees = parse_list_field(request.data.get('referees'))
        additional_details = request.data.get('additional_details', '')
        declaration = request.data.get('declaration', '')
        is_primary = bool(request.data.get('is_primary', True))

        if work_experience and total_years_experience <= 0:
            intervals = []
            for item in work_experience:
                if isinstance(item, dict):
                    sy = item.get('start_year')
                    sm = item.get('start_month', 1) or 1
                    ey = item.get('end_year')
                    em = item.get('end_month', 12) or 12
                    if sy and ey:
                        try:
                            start_idx = int(sy) * 12 + int(sm)
                            end_idx = int(ey) * 12 + int(em)
                            if end_idx >= start_idx:
                                intervals.append((start_idx, end_idx))
                                continue
                        except Exception:
                            pass
                    parsed = parse_date_range_line(item.get('date_range')) if item.get('date_range') else None
                    if parsed:
                        sy, sm, ey, em = parsed
                        start_idx = sy * 12 + sm
                        end_idx = ey * 12 + em
                        if end_idx >= start_idx:
                            intervals.append((start_idx, end_idx))
                else:
                    parsed = parse_date_range_line(item)
                    if parsed:
                        sy, sm, ey, em = parsed
                        start_idx = sy * 12 + sm
                        end_idx = ey * 12 + em
                        if end_idx >= start_idx:
                            intervals.append((start_idx, end_idx))
            if intervals:
                merged = []
                for start, end in sorted(intervals, key=lambda x: x[0]):
                    if not merged or start > merged[-1][1]:
                        merged.append([start, end])
                    else:
                        merged[-1][1] = max(merged[-1][1], end)
                total_months = sum(max(0, end - start + 1) for start, end in merged)
                total_years_experience = round(total_months / 12.0, 2)

        doc = Document()
        doc.add_heading('CARRICULUM VITAE (CV)', level=1)

        doc.add_heading('PERSONAL DETAILS', level=2)
        personal_lines = [
            f"SURNAME: {surname or '-'}",
            f"NAME: {given_names or full_name or '-'}",
            f"SEX: {sex or '-'}",
            f"MARTIAL STATUS: {marital_status or '-'}",
            f"DATE OF BIRTH: {date_of_birth or '-'}",
            f"LANGUAGE: {', '.join(languages) if languages else '-'}",
            f"NATIONALITY: {nationality or '-'}",
            f"CONTACT: {' / '.join([p for p in [phone_number, alt_phone] if p]) or '-'}",
            f"EMAIL: {email or '-'}",
        ]
        for line in personal_lines:
            doc.add_paragraph(line, style='List Bullet')

        doc.add_heading('EDUCATION BACK GROUND', level=2)
        if education:
            for item in education:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph('-')

        doc.add_heading('EXPERIENCE', level=2)
        if work_experience:
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = 'Institution/Organization'
            hdr[1].text = 'Position'
            hdr[2].text = 'Duration'
            hdr[3].text = 'Duties'
            for item in work_experience:
                row = table.add_row().cells
                if isinstance(item, dict):
                    row[0].text = str(item.get('organization') or '')
                    row[1].text = str(item.get('position') or '')
                    row[2].text = str(item.get('date_range') or '')
                    row[3].text = str(item.get('responsibilities') or '')
                else:
                    row[0].text = str(item)
        else:
            doc.add_paragraph('-')

        doc.add_heading('TRAINING AND WORKSHOP', level=2)
        if training_workshop:
            table = doc.add_table(rows=1, cols=2)
            hdr = table.rows[0].cells
            hdr[0].text = 'INSTITUTE'
            hdr[1].text = 'Duties'
            for item in training_workshop:
                row = table.add_row().cells
                if isinstance(item, dict):
                    institute = str(item.get('institution') or '')
                    date_range = str(item.get('date_range') or '')
                    duties = str(item.get('duties') or '')
                    row[0].text = f"{institute} {f'({date_range})' if date_range else ''}".strip()
                    row[1].text = duties
                else:
                    row[0].text = str(item)
        else:
            doc.add_paragraph('-')

        doc.add_heading('SKILLS AND QUALIFICATION', level=2)
        if skills:
            for item in skills:
                doc.add_paragraph(str(item), style='List Bullet')
        else:
            doc.add_paragraph('-')

        if certifications:
            doc.add_paragraph('Certifications:', style='List Bullet')
            for item in certifications:
                doc.add_paragraph(str(item), style='List Bullet')

        doc.add_heading('HOBIES', level=2)
        doc.add_paragraph(hobbies or '-')

        doc.add_heading('REFEREES', level=2)
        if referees:
            table = doc.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = 'Name'
            hdr[1].text = 'Institution/Organization'
            hdr[2].text = 'Title'
            hdr[3].text = 'Contacts'
            for referee in referees:
                row = table.add_row().cells
                if isinstance(referee, dict):
                    row[0].text = str(referee.get('name') or '')
                    row[1].text = str(referee.get('organization') or '')
                    row[2].text = str(referee.get('title') or '')
                    row[3].text = str(referee.get('contact') or '')
                else:
                    parts = [p.strip() for p in str(referee).split('|')]
                    row[0].text = parts[0] if len(parts) > 0 else str(referee)
                    row[1].text = parts[1] if len(parts) > 1 else ''
                    row[2].text = parts[2] if len(parts) > 2 else ''
                    row[3].text = parts[3] if len(parts) > 3 else ''
        else:
            doc.add_paragraph('-')

        doc.add_heading('DECLERATION', level=2)
        if declaration:
            doc.add_paragraph(declaration)
        elif additional_details:
            doc.add_paragraph(additional_details)
        else:
            fallback = f"I’m {full_name or 'the candidate'}, I declare that the information provided is complete and correct to the best of my knowledge."
            doc.add_paragraph(fallback)

        buffer = io.BytesIO()
        doc.save(buffer)
        base_name = (full_name or request.user.username or 'candidate').strip().replace(' ', '_')
        safe_name = ''.join(ch for ch in base_name if ch.isalnum() or ch in ('_', '-'))
        if not safe_name:
            safe_name = 'candidate'
        file_name = f"{safe_name}_cv.docx"
        extracted_text = "\n".join([
            "CARRICULUM VITAE (CV)",
            "PERSONAL DETAILS",
            f"SURNAME: {surname}",
            f"NAME: {given_names or full_name}",
            f"SEX: {sex}",
            f"MARTIAL STATUS: {marital_status}",
            f"DATE OF BIRTH: {date_of_birth}",
            f"LANGUAGE: {', '.join(languages)}",
            f"NATIONALITY: {nationality}",
            f"CONTACT: {' / '.join([p for p in [phone_number, alt_phone] if p])}",
            f"EMAIL: {email}",
            "EDUCATION BACK GROUND",
            f"Education: {' | '.join(map(str, education))}",
            "EXPERIENCE",
            f"Work Experience: {' | '.join(format_work_experience_item(item) for item in work_experience)}",
            "TRAINING AND WORKSHOP",
            f"Training: {' | '.join(format_work_experience_item(item) for item in training_workshop)}",
            "SKILLS AND QUALIFICATION",
            f"Skills: {', '.join(skills)}",
            "HOBIES",
            f"Hobbies: {hobbies}",
            "REFEREES",
            f"Referees: {' | '.join(map(str, referees))}",
            "DECLERATION",
            f"Declaration: {declaration or additional_details}",
        ])

        if is_primary:
            CVUpload.objects.filter(applicant=applicant).update(is_primary=False)

        cv_upload = CVUpload(
            applicant=applicant,
            extracted_text=extracted_text,
            parsed_data={
                'template_source': 'builder',
                'surname': surname,
                'given_names': given_names,
                'sex': sex,
                'marital_status': marital_status,
                'date_of_birth': date_of_birth,
                'nationality': nationality,
                'alt_phone': alt_phone,
                'training_workshop': training_workshop,
                'hobbies': hobbies,
                'declaration': declaration,
                'referees': referees,
                'additional_details': additional_details,
            },
            is_primary=is_primary,
        )
        cv_upload.cv_file.save(file_name, ContentFile(buffer.getvalue()), save=False)
        cv_upload.save()

        ParsedCVData.objects.update_or_create(
            cv_upload=cv_upload,
            defaults={
                'full_name': full_name,
                'email': email,
                'phone_number': phone_number,
                'location': location,
                'summary': summary,
                'skills': skills,
                'education': education,
                'highest_education': highest_education,
                'work_experience': work_experience,
                'total_years_experience': total_years_experience,
                'certifications': certifications,
                'languages': languages,
                'extraction_confidence': 100,
            }
        )

        serializer = self.get_serializer(cv_upload)
        return Response(serializer.data, status=status.HTTP_201_CREATED)


class JobApplicationViewSet(viewsets.ModelViewSet):
    """Job applications management"""
    queryset = JobApplication.objects.all()
    serializer_class = JobApplicationSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        if self.request.user.role == 'applicant':
            applicant = get_object_or_404(Applicant, user=self.request.user)
            return JobApplication.objects.filter(applicant=applicant)
        elif self.request.user.role == 'hr':
            return JobApplication.objects.filter(job__hr__user=self.request.user)
        return JobApplication.objects.none()

    def _rank_job_applications(self, job):
        """Re-rank job applications when new/updated application is submitted."""
        applications = JobApplication.objects.filter(job=job)
        candidates_data = []
        app_map = {}

        for app in applications:
            if not app.cv:
                continue
            parsed = ensure_parsed_data_record(app.cv)
            if not parsed:
                continue
            candidates_data.append({
                'id': str(app.id),
                'skills': parsed.skills,
                'education': parsed.education,
                'highest_education': parsed.highest_education,
                'work_experience': parsed.work_experience,
                'total_years_experience': parsed.total_years_experience,
                'extracted_text': app.cv.extracted_text or '',
            })
            app_map[str(app.id)] = app

        if not candidates_data:
            return

        # Reset previous ranking output for this job before writing fresh ranks.
        applications.update(
            application_score=None,
            application_rank=None,
            skills_score=None,
            experience_score=None,
            education_score=None,
        )

        ranking_engine = RankingEngine({
            'skills_weight': job.screening_criteria.skills_weight,
            'experience_weight': job.screening_criteria.experience_weight,
            'education_weight': job.screening_criteria.education_weight,
        })

        job_data = {
            'required_skills': job.required_skills,
            'preferred_skills': job.preferred_skills,
            'required_education': job.required_education,
            'years_of_experience_min': job.years_of_experience_min,
            'years_of_experience_max': job.years_of_experience_max,
            'job_title': job.job_title,
            'job_description': job.job_description,
        }

        ranked = ranking_engine.rank_candidates(job_data, candidates_data)
        for rank, (candidate, score) in enumerate(ranked, 1):
            app = app_map[candidate['id']]
            app.application_score = score
            app.application_rank = rank
            app.skills_score = candidate['scores']['skills_score']
            app.experience_score = candidate['scores']['experience_score']
            app.education_score = candidate['scores']['education_score']
            app.save(update_fields=[
                'application_score', 'application_rank',
                'skills_score', 'experience_score', 'education_score',
            ])
    
    def create(self, request, *args, **kwargs):
        """Apply for a job (applicant only) with optional explicit CV selection."""
        if request.user.role != 'applicant':
            return Response({'error': 'Only applicants can apply for jobs'}, status=status.HTTP_403_FORBIDDEN)

        applicant = get_object_or_404(Applicant, user=request.user)
        job_id = request.data.get('job')
        cv_id = request.data.get('cv')

        if not job_id:
            return Response({'error': 'Job is required'}, status=status.HTTP_400_BAD_REQUEST)

        job = get_object_or_404(JobPosting, id=job_id, is_active=True)
        selected_cv = None

        if cv_id:
            selected_cv = CVUpload.objects.filter(id=cv_id, applicant=applicant).first()
            if not selected_cv:
                return Response({'error': 'Selected CV not found'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            selected_cv = CVUpload.objects.filter(applicant=applicant, is_primary=True).first()

        if not selected_cv:
            return Response({'error': 'Please upload and set a primary CV first'}, status=status.HTTP_400_BAD_REQUEST)

        existing_application = JobApplication.objects.filter(applicant=applicant, job=job).first()
        if existing_application and not existing_application.reapply_allowed:
            return Response({'error': 'Already applied for this job'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            with transaction.atomic():
                if existing_application and existing_application.reapply_allowed:
                    # Reuse the same application record when HR explicitly allows reapply.
                    existing_application.cv = selected_cv
                    existing_application.status = 'applied'
                    existing_application.application_score = None
                    existing_application.application_rank = None
                    existing_application.skills_score = None
                    existing_application.experience_score = None
                    existing_application.education_score = None
                    existing_application.hr_notes = ''
                    existing_application.reapply_allowed = False
                    existing_application.save(update_fields=[
                        'cv', 'status', 'application_score', 'application_rank',
                        'skills_score', 'experience_score', 'education_score',
                        'hr_notes', 'reapply_allowed', 'updated_date',
                    ])
                    application = existing_application
                else:
                    application = JobApplication.objects.create(
                        applicant=applicant,
                        job=job,
                        cv=selected_cv
                    )
                Notification.objects.create(
                    user=job.hr.user,
                    notification_type='application_status',
                    title='New Application',
                    message=f'{applicant.user.get_full_name()} applied for {job.job_title}',
                    related_job=job,
                    related_application=application
                )
        except IntegrityError:
            return Response({'error': 'Already applied for this job'}, status=status.HTTP_400_BAD_REQUEST)

        serializer = self.get_serializer(application)
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['post'])
    def make_decision(self, request, pk=None):
        """HR makes hiring decision"""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can make decisions'}, status=status.HTTP_403_FORBIDDEN)

        application = self.get_object()
        if application.job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to decide for this application'}, status=status.HTTP_403_FORBIDDEN)
        decision = request.data.get('decision')  # 'hire' or 'reject'
        notes = (request.data.get('notes') or '').strip()
        
        if decision == 'hire':
            application.status = 'hired'
            notify_title = 'Congratulations!'
            notify_message = f'You have been hired for {application.job.job_title} at {application.job.hr.company_name}'
        elif decision == 'reject':
            application.status = 'rejected'
            notify_title = 'Application Update'
            notify_message = f'Your application for {application.job.job_title} has been rejected'
        else:
            return Response({'error': 'Invalid decision'}, status=status.HTTP_400_BAD_REQUEST)

        application.hr_notes = notes
        application.save()

        if notes:
            notify_message = f"{notify_message}\n\nMessage from HR: {notes}"
        
        # Create notification for applicant
        Notification.objects.create(
            user=application.applicant.user,
            notification_type='hiring_decision',
            title=notify_title,
            message=notify_message,
            related_job=application.job,
            related_application=application
        )
        
        serializer = self.get_serializer(application)
        return Response(serializer.data, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def request_help(self, request, pk=None):
        """Applicant requests help from HR for a specific application."""
        if request.user.role != 'applicant':
            return Response({'error': 'Only applicants can request help'}, status=status.HTTP_403_FORBIDDEN)

        application = self.get_object()
        if application.applicant.user_id != request.user.id:
            return Response({'error': 'Not allowed to request help for this application'}, status=status.HTTP_403_FORBIDDEN)

        message = (request.data.get('message') or '').strip()
        if not message:
            return Response({'error': 'Message is required'}, status=status.HTTP_400_BAD_REQUEST)

        title = f'Help Request: {application.job.job_title}'
        body = (
            f'{request.user.get_full_name() or request.user.username} requested help for '
            f'{application.job.job_title}.\n\nMessage:\n{message}'
        )

        Notification.objects.create(
            user=application.job.hr.user,
            notification_type='system',
            title=title,
            message=body,
            related_job=application.job,
            related_application=application
        )

        return Response({'message': 'Help request sent to HR'}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def allow_reapply(self, request, pk=None):
        """HR allows candidate to re-apply for the same job once."""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can allow reapply'}, status=status.HTTP_403_FORBIDDEN)

        application = self.get_object()
        if application.job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to modify this application'}, status=status.HTTP_403_FORBIDDEN)

        application.reapply_allowed = True
        application.save(update_fields=['reapply_allowed'])

        Notification.objects.create(
            user=application.applicant.user,
            notification_type='application_status',
            title=f'Reapply Enabled: {application.job.job_title}',
            message='HR has allowed you to apply again for this job.',
            related_job=application.job,
            related_application=application
        )

        return Response({'message': 'Candidate can now reapply for this job'}, status=status.HTTP_200_OK)


class NotificationViewSet(viewsets.ModelViewSet):
    """Notification management"""
    queryset = Notification.objects.all()
    serializer_class = NotificationSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        return Notification.objects.filter(user=self.request.user)
    
    @action(detail=False, methods=['get'])
    def unread(self, request):
        """Get unread notifications"""
        notifications = Notification.objects.filter(user=request.user, is_read=False)
        serializer = self.get_serializer(notifications, many=True)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def mark_as_read(self, request, pk=None):
        """Mark notification as read"""
        notification = self.get_object()
        notification.is_read = True
        notification.read_at = timezone.now()
        notification.save()
        return Response({'message': 'Marked as read'}, status=status.HTTP_200_OK)

    @action(detail=True, methods=['post'])
    def reply_help(self, request, pk=None):
        """HR replies to candidate help request notification."""
        if request.user.role != 'hr':
            return Response({'error': 'Only HR users can reply to help requests'}, status=status.HTTP_403_FORBIDDEN)

        notification = self.get_object()
        if notification.user_id != request.user.id:
            return Response({'error': 'Not allowed to reply to this notification'}, status=status.HTTP_403_FORBIDDEN)

        message = (request.data.get('message') or '').strip()
        if not message:
            return Response({'error': 'Reply message is required'}, status=status.HTTP_400_BAD_REQUEST)

        application = notification.related_application
        if not application:
            return Response({'error': 'This help request is not linked to an application'}, status=status.HTTP_400_BAD_REQUEST)

        if application.job.hr.user_id != request.user.id:
            return Response({'error': 'Not allowed to reply for this application'}, status=status.HTTP_403_FORBIDDEN)

        Notification.objects.create(
            user=application.applicant.user,
            notification_type='system',
            title=f'HR Reply: {application.job.job_title}',
            message=f'Reply from HR:\n{message}',
            related_job=application.job,
            related_application=application
        )

        notification.is_read = True
        notification.read_at = timezone.now()
        notification.save(update_fields=['is_read', 'read_at'])

        return Response({'message': 'Reply sent to candidate'}, status=status.HTTP_200_OK)


